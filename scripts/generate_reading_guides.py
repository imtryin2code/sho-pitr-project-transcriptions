import csv
import os
import xml.sax.saxutils as saxutils
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ReportLab Imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

def safe_reportlab_text(text):
    """
    Safely encodes brackets so ReportLab renders them as literal text 
    instead of parsing or stripping them as structural HTML tags.
    """
    if not text:
        return ""
    # Replace amp first so we don't accidentally double-encode later replacements
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    return text

def generate_guides():
    csv_path = 'metadata/master_transcription_list.csv'
    md_output_base = 'exports/markdown'
    pdf_output_base = 'exports/pdfs'
    word_output_base = 'exports/word-docs'
    font_path = "scripts/DejaVuSans.ttf" 

    if not os.path.exists(csv_path):
        print("Error: Master CSV not found.")
        return

    os.makedirs(md_output_base, exist_ok=True)
    os.makedirs(pdf_output_base, exist_ok=True)
    os.makedirs(word_output_base, exist_ok=True)

    # Register the Unicode Font
    font_name = "Helvetica"
    font_bold_name = "Helvetica-Bold"
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
        font_name = "DejaVuSans"
        font_bold_name = "DejaVuSans" 
    else:
        print(f"Warning: Font not found at {font_path}, falling back to Helvetica.")

    stories = {}
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            story_id = row['ID']
            if story_id not in stories:
                stories[story_id] = []
            stories[story_id].append(row)

    for story_id, rows in stories.items():
        print(f"Processing {story_id}...")

        # --- 1. MARKDOWN GENERATION ---
        md_path = os.path.join(md_output_base, f"{story_id}_Reading_Guide.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# Reading Guide: {story_id}\n\n")
            f.write("| Time | Speaker | Text | Research Notes |\n")
            f.write("| :--- | :--- | :--- | :--- |\n")
            for raw_row in rows:
                # Isolate row data completely
                row = raw_row.copy()
                md_text = row['Text'] if row['Text'] else ""
                
                if any(name in row['Speaker'] for name in ["Joe", "Peter"]):
                    md_text = f"**{md_text}**"
                
                note_display = ""
                if row.get('Notes_Text') and row.get('Notes_Tier'):
                    note_display = f"*{row['Notes_Tier']}*: {row['Notes_Text']}"
                f.write(f"| {row['Time']} | {row['Speaker']} | {md_text} | {note_display} |\n")

        # --- 2. REPORTLAB PDF GENERATION ---
        pdf_path = os.path.join(pdf_output_base, f"{story_id}_Reading_Guide.pdf")
        
        doc = SimpleDocTemplate(
            pdf_path, 
            pagesize=letter,
            leftMargin=36, rightMargin=36, 
            topMargin=36, bottomMargin=36
        )
        
        # Setup Styles
        style_header = ParagraphStyle('HeaderStyle', fontName='Helvetica-Bold', fontSize=10, textColor=colors.white, alignment=1)
        style_time = ParagraphStyle('TimeStyle', fontName='Helvetica', fontSize=9, alignment=1)
        style_speaker = ParagraphStyle('SpeakerStyle', fontName='Helvetica', fontSize=9)
        style_speaker_bold = ParagraphStyle('SpeakerStyleBold', fontName='Helvetica-Bold', fontSize=9)
        style_text = ParagraphStyle('TextStyle', fontName=font_name, fontSize=9, leading=12)
        style_text_bold = ParagraphStyle('TextStyleBold', fontName=font_bold_name, fontSize=9, leading=12)
        style_notes = ParagraphStyle('NotesStyle', fontName=font_name, fontSize=8, leading=10, textColor=colors.HexColor('#444444'))

        pdf_elements = []
        title_style = ParagraphStyle('TitleStyle', fontName='Helvetica-Bold', fontSize=14, alignment=1, spaceAfter=15)
        pdf_elements.append(Paragraph(f"Joe Peter Project: {story_id}", title_style))
        
        col_widths = [45, 75, 240, 180] 
        
        table_data = [[
            Paragraph("Time", style_header),
            Paragraph("Speaker", style_header),
            Paragraph("Transcription", style_header),
            Paragraph("Research Notes", style_header)
        ]]

        for raw_row in rows:
            # Isolate row data completely
            row = raw_row.copy()
            is_joe = any(name in row['Speaker'] for name in ["Joe", "Peter"])
            
            # Protect the literal tokens and clean strings cleanly
            clean_text = safe_reportlab_text(row['Text'])
            clean_note_tier = safe_reportlab_text(row['Notes_Tier'])
            clean_note_text = safe_reportlab_text(row['Notes_Text'])
            clean_speaker = safe_reportlab_text(row['Speaker'])
            
            # Apply explicit bold styling ONLY via clean PDF tags, ignoring raw Markdown asterisks
            if is_joe:
                txt_p = Paragraph(f"<b>{clean_text}</b>", style_text_bold)
                speaker_p = Paragraph(f"<b>{clean_speaker}</b>", style_speaker_bold)
            else:
                txt_p = Paragraph(clean_text, style_text)
                speaker_p = Paragraph(clean_speaker, style_speaker)
            
            note_content = f"<b>[{clean_note_tier}]</b> {clean_note_text}" if clean_note_text else ""
            note_p = Paragraph(note_content, style_notes)
            time_p = Paragraph(row['Time'], style_time)

            table_data.append([time_p, speaker_p, txt_p, note_p])

        # Construct Table
        guide_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        guide_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8c1b1b')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        pdf_elements.append(guide_table)
        doc.build(pdf_elements)

        # --- 3. WORD DOCUMENT GENERATION ---
        doc_word = Document()
        doc_title = doc_word.add_heading(f'Reading Guide: {story_id}', 0)
        doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table_word = doc_word.add_table(rows=1, cols=4)
        table_word.style = 'Table Grid'
        hdr_cells = table_word.rows[0].cells
        hdr_cells[0].text = 'Time'
        hdr_cells[1].text = 'Speaker'
        hdr_cells[2].text = 'Transcription'
        hdr_cells[3].text = 'Research Category & Notes'

        for raw_row in rows:
            row = raw_row.copy()
            row_cells = table_word.add_row().cells
            row_cells[0].text = row['Time']
            row_cells[1].text = row['Speaker']
            
            p_trans = row_cells[2].paragraphs[0]
            run_trans = p_trans.add_run(row['Text'])
            if any(name in row['Speaker'] for name in ["Joe", "Peter"]):
                run_trans.bold = True
                
            if row.get('Notes_Text'):
                p_notes = row_cells[3].paragraphs[0]
                run_notes = p_notes.add_run(f"[{row['Notes_Tier']}] {row['Notes_Text']}")
                run_notes.italic = True

        doc_path = os.path.join(word_output_base, f"{story_id}_Reading_Guide.docx")
        doc_word.save(doc_path)
        
    print(f"\nSuccess! Exported completely isolated PDF, MD, and DOCX files.")

if __name__ == "__main__":
    generate_guides()